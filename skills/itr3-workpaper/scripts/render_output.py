#!/usr/bin/env python3
"""
Deterministic renderer for the itr3-workpaper skill.

Takes a workpaper JSON (sections + reconciliations) and renders:
  1. <slug>_itr3_workpaper.docx
  2. <slug>_itr3_workpaper.html
Optionally validates an e-filing JSON against the official ITR-3 schema and writes:
  3. <slug>_itr3_efiling_upload.json

Usage:
    python3 render_output.py <workpaper.json> <output_dir> [efiling.json]

No tax computation logic lives here — layout + schema validation only.
"""
import sys
import os
import json
import html as htmlmod
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PREPARED_BY_NAME = "CA Shailesh S Wadhawaniya"
PREPARED_BY_PHONE = "+91 79900 46540"
PREPARED_BY_URL = "https://shaileshai.qzz.io/"

STATUS_META = {
    "filled": {
        "label": "",
        "color_docx": None,
        "css_class": "status-filled",
    },
    "missing": {
        "label": "[MISSING — CONFIRM WITH CLIENT]",
        "color_docx": RGBColor(0xC0, 0x00, 0x00),
        "css_class": "status-missing",
    },
    "judgment": {
        "label": "[JUDGMENT CALL — CA TO CONFIRM]",
        "color_docx": RGBColor(0xB8, 0x6A, 0x00),
        "css_class": "status-judgment",
    },
    "not_applicable": {
        "label": "[N/A]",
        "color_docx": RGBColor(0x70, 0x70, 0x70),
        "css_class": "status-na",
    },
}

RECON_META = {
    "matched": {"label": "MATCHED", "color": RGBColor(0x1A, 0x7A, 0x3C)},
    "mismatch": {"label": "MISMATCH", "color": RGBColor(0xC0, 0x00, 0x00)},
    "missing": {"label": "MISSING", "color": RGBColor(0xB8, 0x6A, 0x00)},
}

# Canonical section checklist — warn if any omitted from workpaper sections[]
CANONICAL_SECTIONS = [
    "Cover",
    "PartA_GEN1",
    "PartA_GEN2",
    "PARTA_BS",
    "ManufacturingAccount",
    "TradingAccount",
    "PARTA_PL",
    "PARTA_OI",
    "PARTA_QD",
    "ScheduleS",
    "ScheduleHP",
    "ITR3ScheduleBP",
    "ScheduleDPM",
    "ScheduleDOA",
    "ScheduleDEP",
    "ScheduleDCG",
    "ScheduleESR",
    "ScheduleCG",
    "Schedule112A",
    "Schedule115AD",
    "ScheduleVDA",
    "ScheduleOS",
    "ScheduleCYLA",
    "ScheduleBFLA",
    "ScheduleCFL",
    "ScheduleICDS",
    "ScheduleUD",
    "ScheduleVIA_and_80s",
    "ScheduleAMT",
    "ScheduleAMTC",
    "ScheduleSI",
    "ScheduleSPI",
    "ScheduleIF",
    "ScheduleEI",
    "SchedulePTI",
    "ScheduleIT",
    "ScheduleTDS1",
    "ScheduleTDS2",
    "ScheduleTDS3",
    "ScheduleTCS",
    "ScheduleFSI",
    "ScheduleTR",
    "ScheduleFA",
    "ScheduleAL",
    "ScheduleGST",
    "ScheduleESOP",
    "Schedule5A",
    "PartB-TI",
    "PartB_TTI",
    "Verification",
]


def slugify(name):
    return "".join(c if c.isalnum() else "_" for c in name).strip("_")[:60] or "client"


def cell_set_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_kv_table(doc, mapping):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, val in mapping.items():
        row = table.add_row().cells
        row[0].text = str(key).replace("_", " ").title()
        row[1].text = "" if val is None else str(val)
    return table


def add_data_table(doc, rows):
    if not rows:
        return
    if isinstance(rows[0], dict):
        headers = list(rows[0].keys())
        t = doc.add_table(rows=0, cols=len(headers))
        t.style = "Light Grid Accent 1"
        hdr = t.add_row().cells
        for i, hcol in enumerate(headers):
            hdr[i].text = str(hcol)
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
            cell_set_shading(hdr[i], "D9E1F2")
        for row in rows:
            cells = t.add_row().cells
            for i, hcol in enumerate(headers):
                cells[i].text = "" if row.get(hcol) is None else str(row.get(hcol))
    else:
        t = doc.add_table(rows=0, cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = "" if v is None else str(v)


def status_counts(sections):
    counts = {"filled": 0, "missing": 0, "judgment": 0, "not_applicable": 0}
    for s in sections:
        st = s.get("status", "filled")
        counts[st] = counts.get(st, 0) + 1
    return counts


def omitted_sections(sections):
    present = set(s.get("id") for s in sections if s.get("id"))
    expected = set(CANONICAL_SECTIONS)
    omitted = [x for x in CANONICAL_SECTIONS if x not in present]
    return omitted


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def render_docx(data, out_path):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    client = data.get("client", {})
    ay = data.get("assessment_year", client.get("assessment_year", "2026-27"))

    title = doc.add_heading(f"ITR-3 Workpaper — AY {ay}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Draft prepared {datetime.now().strftime('%d-%b-%Y')} — for CA / taxpayer review before filing"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    banner = doc.add_paragraph()
    br = banner.add_run(
        "DRAFT ONLY — does not file, sign, or replace professional judgment. "
        "Schema-valid e-filing JSON (if produced) is not a portal-acceptance guarantee."
    )
    br.bold = True
    br.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_heading("Client Details", level=1)
    add_kv_table(doc, client)

    sections = data.get("sections", [])
    counts = status_counts(sections)

    doc.add_heading("Review Summary", level=1)
    sp = doc.add_paragraph()
    sp.add_run(
        f"{len(sections)} sections drafted — "
        f"{counts.get('filled', 0)} filled, "
        f"{counts.get('judgment', 0)} flagged for CA judgment, "
        f"{counts.get('missing', 0)} missing data, "
        f"{counts.get('not_applicable', 0)} not applicable."
    ).bold = True

    missing_ids = [s["id"] for s in sections if s.get("status") == "missing" and s.get("id")]
    judgment_ids = [s["id"] for s in sections if s.get("status") == "judgment" and s.get("id")]
    if missing_ids:
        doc.add_paragraph(f"Sections with missing data: {', '.join(missing_ids)}")
    if judgment_ids:
        doc.add_paragraph(f"Sections needing CA judgment: {', '.join(judgment_ids)}")

    omitted = omitted_sections(sections)
    if omitted:
        warn = doc.add_paragraph()
        warn.add_run(
            f"⚠ {len(omitted)} canonical ITR-3 section(s) do not appear in this draft at all "
            f"(not even as N/A): {', '.join(omitted)}. Confirm whether deliberately skipped or missed."
        ).font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Reconciliation dashboard
    recons = data.get("reconciliations", [])
    if recons:
        doc.add_heading("Reconciliation Dashboard", level=1)
        rcounts = {"matched": 0, "mismatch": 0, "missing": 0}
        for r in recons:
            rcounts[r.get("status", "missing")] = rcounts.get(r.get("status", "missing"), 0) + 1
        doc.add_paragraph(
            f"{rcounts.get('matched', 0)} matched · "
            f"{rcounts.get('mismatch', 0)} mismatch · "
            f"{rcounts.get('missing', 0)} missing source"
        )
        add_data_table(
            doc,
            [
                {
                    "Item": r.get("item", ""),
                    "Status": r.get("status", ""),
                    "Form16/Books": r.get("form16", r.get("books", "")),
                    "AIS/26AS": r.get("ais", r.get("form26as", "")),
                    "Workpaper": r.get("workpaper_amount", ""),
                    "Notes": r.get("notes", ""),
                }
                for r in recons
            ],
        )

    doc.add_page_break()
    doc.add_heading("Schedule-by-Schedule Workpaper", level=1)

    for s in sections:
        status = s.get("status", "filled")
        meta = STATUS_META.get(status, STATUS_META["filled"])
        sid = s.get("id", "?")
        doc.add_heading(f"{sid} — {s.get('title', '')}", level=2)

        if meta["label"]:
            flag_p = doc.add_paragraph()
            run = flag_p.add_run(meta["label"])
            run.bold = True
            if meta["color_docx"]:
                run.font.color.rgb = meta["color_docx"]

        answer = s.get("answer")
        if answer:
            doc.add_paragraph(str(answer))

        rows = s.get("table")
        if rows:
            add_data_table(doc, rows)

        src = s.get("source_refs")
        if src:
            p = doc.add_paragraph()
            p.add_run("Sources: " + ("; ".join(src) if isinstance(src, list) else str(src))).italic = True

        notes = s.get("notes")
        if notes:
            np_ = doc.add_paragraph()
            nrun = np_.add_run(f"Note: {notes}")
            nrun.italic = True
            nrun.font.size = Pt(9.5)

    doc.add_page_break()
    credit = doc.add_paragraph()
    credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = credit.add_run(
        f"Workpaper prepared by {PREPARED_BY_NAME}  |  {PREPARED_BY_PHONE}  |  {PREPARED_BY_URL}"
    )
    crun.italic = True
    crun.font.size = Pt(9)
    crun.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.save(out_path)


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>ITR-3 Workpaper - {client_name}</title>
<style>
  :root {{
    --bg: #ffffff; --fg: #1a1a1a; --muted: #666; --border: #ddd;
    --header-bg: #0f3d2e; --header-fg: #ffffff;
    --filled: #1a7a3c; --missing: #c00000; --judgment: #b86a00; --na: #707070;
    --card-bg: #f7f8fa; --mismatch: #c00000; --matched: #1a7a3c;
  }}
  @media (prefers-color-scheme: dark) {{
    :root {{ --bg:#12151a; --fg:#e8e8e8; --muted:#9aa0a6; --border:#333;
      --header-bg:#0a2a20; --header-fg:#eef4ff; --card-bg:#1b1f26; }}
  }}
  * {{ box-sizing: border-box; }}
  body {{ margin:0; background:var(--bg); color:var(--fg); font-family: "Segoe UI", Arial, sans-serif; line-height:1.5; }}
  header.page {{ background:var(--header-bg); color:var(--header-fg); padding:2rem 1.5rem; }}
  header.page h1 {{ margin:0 0 .25rem 0; font-size:1.5rem; }}
  header.page p {{ margin:0; opacity:.85; font-size:.9rem; }}
  .wrap {{ max-width: 980px; margin: 0 auto; padding: 1.5rem; }}
  .banner {{ border:1px solid var(--missing); background:rgba(192,0,0,.08); border-radius:8px; padding:.75rem 1rem; margin-bottom:1.5rem; color:var(--missing); font-weight:600; font-size:.9rem; }}
  .client-card, .summary-card {{ background:var(--card-bg); border:1px solid var(--border); border-radius:8px; padding:1rem 1.25rem; margin-bottom:1.5rem; }}
  .client-card table {{ width:100%; border-collapse:collapse; }}
  .client-card td {{ padding:.25rem .5rem; vertical-align:top; }}
  .client-card td:first-child {{ color:var(--muted); width:220px; }}
  .summary-pills {{ display:flex; gap:.75rem; flex-wrap:wrap; margin-top:.5rem; }}
  .pill {{ border-radius:999px; padding:.25rem .75rem; font-size:.85rem; font-weight:600; }}
  .pill.filled {{ background:rgba(26,122,60,.15); color:var(--filled); }}
  .pill.missing {{ background:rgba(192,0,0,.15); color:var(--missing); }}
  .pill.judgment {{ background:rgba(184,106,0,.15); color:var(--judgment); }}
  .pill.na {{ background:rgba(112,112,112,.15); color:var(--na); }}
  .omitted-warning {{ border:1px solid var(--missing); background:rgba(192,0,0,.08); border-radius:8px; padding:.75rem 1rem; margin-bottom:1.5rem; color:var(--missing); font-size:.9rem; }}
  nav.toc {{ columns:2; column-gap:1rem; font-size:.85rem; margin-bottom:2rem; }}
  nav.toc a {{ color:var(--fg); text-decoration:none; display:block; padding:.15rem 0; }}
  nav.toc a:hover {{ text-decoration:underline; }}
  section.clause {{ border:1px solid var(--border); border-radius:8px; padding:1.25rem 1.5rem; margin-bottom:1.25rem; scroll-margin-top:1rem; }}
  section.clause h2 {{ margin:0 0 .5rem 0; font-size:1.05rem; }}
  .badge {{ display:inline-block; font-weight:700; font-size:.8rem; padding:.15rem .6rem; border-radius:4px; margin-bottom:.5rem; }}
  .status-missing .badge {{ color:#fff; background:var(--missing); }}
  .status-judgment .badge {{ color:#fff; background:var(--judgment); }}
  .status-na .badge {{ color:#fff; background:var(--na); }}
  .answer {{ white-space:pre-wrap; }}
  table.data {{ width:100%; border-collapse:collapse; margin-top:.75rem; font-size:.9rem; }}
  table.data th, table.data td {{ border:1px solid var(--border); padding:.4rem .6rem; text-align:left; }}
  table.data th {{ background:var(--card-bg); }}
  .note {{ color:var(--muted); font-size:.85rem; margin-top:.6rem; font-style:italic; }}
  .src {{ color:var(--muted); font-size:.85rem; margin-top:.4rem; }}
  .recon-mismatch {{ color:var(--mismatch); font-weight:700; }}
  .recon-matched {{ color:var(--matched); font-weight:700; }}
  footer {{ text-align:center; color:var(--muted); font-size:.8rem; padding:2rem 1rem; }}
</style>
</head>
<body>
<header class="page">
  <h1>ITR-3 Workpaper &mdash; AY {ay}</h1>
  <p>Draft prepared {gen_date} &middot; for CA / taxpayer review before filing &middot; not a final return</p>
</header>
<div class="wrap">
  <div class="banner">DRAFT ONLY &mdash; does not file, sign, or replace professional judgment. Schema validation &ne; portal acceptance.</div>
  <div class="client-card">
    <table>{client_rows}</table>
  </div>
  <div class="summary-card">
    <strong>{n_sections} sections drafted</strong>
    <div class="summary-pills">
      <span class="pill filled">{n_filled} filled</span>
      <span class="pill judgment">{n_judgment} judgment calls</span>
      <span class="pill missing">{n_missing} missing data</span>
      <span class="pill na">{n_na} not applicable</span>
    </div>
  </div>
  {omitted_html}
  {recon_html}
  <nav class="toc">{toc_html}</nav>
  {sections_html}
</div>
<footer>
  Generated by itr3-workpaper skill &middot; verify every flagged item before filing<br>
  Workpaper prepared by {prepared_by_name} &middot; {prepared_by_phone} &middot;
  <a href="{prepared_by_url}" style="color:inherit;">{prepared_by_url}</a>
</footer>
</body>
</html>
"""


def esc(v):
    return htmlmod.escape(str(v)) if v is not None else ""


def table_html(rows):
    if not rows:
        return ""
    if isinstance(rows[0], dict):
        headers = list(rows[0].keys())
        thead = "".join(f"<th>{esc(h)}</th>" for h in headers)
        trs = "".join(
            "<tr>" + "".join(f"<td>{esc(row.get(h))}</td>" for h in headers) + "</tr>"
            for row in rows
        )
    else:
        thead = ""
        trs = "".join(
            "<tr>" + "".join(f"<td>{esc(v)}</td>" for v in row) + "</tr>" for row in rows
        )
    return f'<table class="data"><thead><tr>{thead}</tr></thead><tbody>{trs}</tbody></table>'


def render_html(data, out_path):
    client = data.get("client", {})
    ay = data.get("assessment_year", client.get("assessment_year", "2026-27"))
    sections = data.get("sections", [])

    client_rows = "".join(
        f"<tr><td>{esc(k.replace('_', ' ').title())}</td><td>{esc(v)}</td></tr>"
        for k, v in client.items()
    )

    counts = status_counts(sections)
    omitted = omitted_sections(sections)
    omitted_html = ""
    if omitted:
        omitted_html = (
            f'<div class="omitted-warning">&#9888; {len(omitted)} canonical ITR-3 section(s) do not '
            f"appear in this draft at all (not even as N/A): {esc(', '.join(omitted))}. "
            f"Confirm whether deliberately skipped or missed.</div>"
        )

    recons = data.get("reconciliations", [])
    recon_html = ""
    if recons:
        rrows = []
        for r in recons:
            st = r.get("status", "")
            cls = "recon-mismatch" if st == "mismatch" else ("recon-matched" if st == "matched" else "")
            rrows.append(
                {
                    "Item": r.get("item", ""),
                    "Status": st,
                    "Form16/Books": r.get("form16", r.get("books", "")),
                    "AIS/26AS": r.get("ais", r.get("form26as", "")),
                    "Workpaper": r.get("workpaper_amount", ""),
                    "Notes": r.get("notes", ""),
                }
            )
        recon_html = (
            '<div class="summary-card"><strong>Reconciliation Dashboard</strong>'
            + table_html(rrows)
            + "</div>"
        )

    toc_html = "".join(
        f'<a href="#sec-{esc(s.get("id", i))}">{esc(s.get("id", ""))} &mdash; {esc(s.get("title", ""))}</a>'
        for i, s in enumerate(sections)
    )

    blocks = []
    for i, s in enumerate(sections):
        status = s.get("status", "filled")
        meta = STATUS_META.get(status, STATUS_META["filled"])
        badge = f'<div class="badge">{esc(meta["label"])}</div>' if meta["label"] else ""
        answer_html = f'<div class="answer">{esc(s.get("answer", ""))}</div>' if s.get("answer") else ""
        note_html = f'<div class="note">Note: {esc(s.get("notes"))}</div>' if s.get("notes") else ""
        src = s.get("source_refs")
        src_html = ""
        if src:
            src_html = (
                f'<div class="src">Sources: {esc("; ".join(src) if isinstance(src, list) else src)}</div>'
            )
        sid = s.get("id", str(i))
        blocks.append(
            f'<section class="clause {meta["css_class"]}" id="sec-{esc(sid)}">'
            f'<h2>{esc(sid)} &mdash; {esc(s.get("title", ""))}</h2>'
            f"{badge}{answer_html}{table_html(s.get('table'))}{src_html}{note_html}"
            f"</section>"
        )

    html_out = HTML_TEMPLATE.format(
        client_name=esc(client.get("name", "Client")),
        ay=esc(ay),
        gen_date=datetime.now().strftime("%d-%b-%Y"),
        client_rows=client_rows,
        n_sections=len(sections),
        n_filled=counts.get("filled", 0),
        n_judgment=counts.get("judgment", 0),
        n_missing=counts.get("missing", 0),
        n_na=counts.get("not_applicable", 0),
        omitted_html=omitted_html,
        recon_html=recon_html,
        toc_html=toc_html,
        sections_html="\n".join(blocks),
        prepared_by_name=esc(PREPARED_BY_NAME),
        prepared_by_phone=esc(PREPARED_BY_PHONE),
        prepared_by_url=esc(PREPARED_BY_URL),
    )
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html_out)


# ---------------------------------------------------------------------------
# E-filing validation
# ---------------------------------------------------------------------------

def validate_and_write_efiling(efiling_path, out_dir, slug):
    try:
        import jsonschema
    except ImportError:
        print(
            "Skipping e-filing JSON validation: install jsonschema (pip install jsonschema). "
            "docx/html outputs are unaffected.",
            file=sys.stderr,
        )
        return False

    with open(efiling_path, "r", encoding="utf-8") as f:
        efiling_data = json.load(f)

    if "ITR" not in efiling_data or "ITR3" not in efiling_data.get("ITR", {}):
        print(
            'Cannot validate e-filing JSON: expected top-level {"ITR": {"ITR3": {...}}} — '
            f"found top-level keys: {list(efiling_data.keys())}",
            file=sys.stderr,
        )
        return False

    schema_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..",
        "references",
        "schema",
        "schema_ITR3.json",
    )
    with open(schema_path, "r", encoding="utf-8") as f:
        schema = json.load(f)

    validator = jsonschema.Draft4Validator(schema)
    errors = sorted(validator.iter_errors(efiling_data), key=lambda e: [str(p) for p in e.path])

    if errors:
        print(
            f"E-FILING JSON FAILED SCHEMA VALIDATION ({len(errors)} error(s)) against schema_ITR3.json:",
            file=sys.stderr,
        )
        for e in errors[:60]:
            path = "/".join(str(p) for p in e.path) or "(root)"
            print(f"  - {path}: {e.message}", file=sys.stderr)
        if len(errors) > 60:
            print(f"  ... and {len(errors) - 60} more error(s)", file=sys.stderr)
        print(
            "Fix the e-filing JSON and re-run. The docx/html above were still produced. "
            "Default policy: invalid e-filing file is NOT written.",
            file=sys.stderr,
        )
        return False

    efiling_out_path = os.path.join(out_dir, f"{slug}_itr3_efiling_upload.json")
    with open(efiling_out_path, "w", encoding="utf-8") as f:
        json.dump(efiling_data, f, indent=2, ensure_ascii=False)
        f.write("\n")

    print(f"Written: {efiling_out_path}  (validated against schema_ITR3.json, ITR-3 Main V1.1)")
    print(
        "Note: schema conformance only — not portal Category A acceptance, not a filed return. "
        "Complete remaining fields in official utility / firm software."
    )
    return True


def main():
    if len(sys.argv) not in (3, 4):
        print(
            "Usage: python3 render_output.py <workpaper.json> <output_dir> [efiling.json]",
            file=sys.stderr,
        )
        sys.exit(1)

    in_path, out_dir = sys.argv[1], sys.argv[2]
    efiling_path = sys.argv[3] if len(sys.argv) == 4 else None

    with open(in_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    os.makedirs(out_dir, exist_ok=True)
    slug = slugify(data.get("client", {}).get("name", "client"))
    docx_path = os.path.join(out_dir, f"{slug}_itr3_workpaper.docx")
    html_path = os.path.join(out_dir, f"{slug}_itr3_workpaper.html")

    render_docx(data, docx_path)
    render_html(data, html_path)

    print(f"Written: {docx_path}")
    print(f"Written: {html_path}")

    if efiling_path:
        if not validate_and_write_efiling(efiling_path, out_dir, slug):
            sys.exit(1)


if __name__ == "__main__":
    main()
